import docx
import re
import sys, os

def usage():
    None



def match(a): 
    rok = re.search(r'\(((\d)|(\*)).*?\)', a) 
    meno = re.search(r'^.*?((\t)|(\())',a) 
    skladba = re.search(r'\t.*$',a) 
    if rok != None: 
        rok = rok.group()
    else:
        rok = ''
    if meno != None: 
        meno = meno.group()[:-1]
    else:
        meno = '' 
    if skladba != None: 
        skladba = re.sub(r'\t','',skladba.group())
    else:
        skladba = '' 
    return [meno, rok, skladba] 




if __name__ == "__main__":
    if len(sys.argv) < 2:
        usage()
    doc = docx.Document(sys.argv[1])

    program = []
    for i in range(len(doc.paragraphs)):
        meno, rok, skladba = match(doc.paragraphs[i].text)
        if not any([meno,rok, skladba]):
            continue
        if not any([meno, rok]):
            program[-1]["skladby"].append(skladba)
            continue
        program.append({"meno": meno,
        "rok": rok,
        "skladby": [skladba]
        })


    for item in program:
        print('{{< skladatel meno="%s" rok="%s">}}' % (item["meno"], item["rok"]))
        print("{{< skladby >}}")
        for item2 in item["skladby"]:
            print("**%s** {{< br >}}" %item2)
        print('{{< /skladby >}} ')
        print()
     
